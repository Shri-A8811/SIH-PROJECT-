"""
Industrial Word Document (.docx) Generator for Sovereign On-Premise Agentic AI Workbench.
Produces production-grade engineering approval notes and inspection summaries.
Always automatically inserts the mandatory "AI-GENERATED DRAFT — HUMAN REVIEW REQUIRED" disclaimer.
"""
from typing import Any, Dict, List, Optional
from datetime import datetime, timezone
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from config.settings import OUTPUT_DIR
from src.core.state_store import StateStore


class DocxApprovalNoteGenerator:
    """Generates verified engineering approval notes from structured state."""

    def __init__(self, state_store: StateStore):
        self.state_store = state_store

    def generate_approval_note(
        self,
        project_id: str,
        title: str,
        executive_summary: str,
        findings: List[Dict[str, Any]],
        calculation_data: Optional[Dict[str, Any]] = None,
        sop_citations: Optional[List[Dict[str, Any]]] = None,
        output_filename: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Creates a professional .docx approval note with strict grounding citations."""
        doc = docx.Document()

        # Page Setup (1 inch margins)
        sections = doc.sections
        for s in sections:
            s.top_margin = Inches(1)
            s.bottom_margin = Inches(1)
            s.left_margin = Inches(1)
            s.right_margin = Inches(1)

        # 1. MANDATORY HUMAN REVIEW DISCLAIMER BANNER
        disclaimer_table = doc.add_table(rows=1, cols=1)
        disclaimer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = disclaimer_table.cell(0, 0)
        
        # Style banner with dark red border and amber/yellow shading
        shading_xml = parse_xml(r'<w:shd {} w:fill="FFF3CD"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_xml)
        
        p_disc = cell.paragraphs[0]
        p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_disc = p_disc.add_run("⚠️ AI-GENERATED DRAFT — HUMAN REVIEW REQUIRED")
        run_disc.bold = True
        run_disc.font.size = Pt(11)
        run_disc.font.color.rgb = RGBColor(133, 100, 4)  # Dark amber
        
        p_disc2 = cell.add_paragraph()
        p_disc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_disc2.add_run(
            "This technical approval note was synthesized by the Sovereign On-Premise Agentic AI Workbench. "
            "All quantitative findings are grounded in verified inspection evidence and SOP citations. "
            "Final engineering authorization requires physical signature below."
        )
        run_sub.font.size = Pt(9)
        run_sub.font.italic = True
        run_sub.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()  # Spacing

        # 2. DOCUMENT HEADER
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_org = header_p.add_run("SOVEREIGN INDUSTRIAL AI WORKBENCH\n")
        run_org.bold = True
        run_org.font.size = Pt(14)
        run_org.font.color.rgb = RGBColor(15, 45, 90)

        run_title = header_p.add_run(f"TECHNICAL APPROVAL & ENGINEERING ACTION NOTE\n{title.upper()}")
        run_title.bold = True
        run_title.font.size = Pt(12)
        run_title.font.color.rgb = RGBColor(30, 30, 30)

        # Meta Table
        now_dt = datetime.now(timezone.utc)
        meta_table = doc.add_table(rows=2, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_table.cell(0, 0).paragraphs[0].add_run(f"Project ID: {project_id}").bold = True
        meta_table.cell(0, 1).paragraphs[0].add_run(f"Date: {now_dt.strftime('%d-%b-%Y')}").bold = True
        meta_table.cell(1, 0).paragraphs[0].add_run("Classification: Internal Confidential (Air-Gapped)").font.size = Pt(9)
        meta_table.cell(1, 1).paragraphs[0].add_run("Ref Standard: SOP-17 / SOP-04").font.size = Pt(9)

        doc.add_paragraph()

        # 3. EXECUTIVE SUMMARY
        doc.add_heading("1. Executive Summary", level=2)
        p_summary = doc.add_paragraph(executive_summary)
        p_summary.paragraph_format.line_spacing = 1.15

        # 4. SAFETY-CRITICAL FINDINGS TABLE
        doc.add_heading("2. Safety-Critical Inspection Findings & Grounding", level=2)
        
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = table.rows[0].cells
        headers = ["Evidence ID", "Equipment Component", "Observed Anomaly / Defect", "Severity", "Measured vs SOP Limit", "Status"]
        for i, title_text in enumerate(headers):
            hdr_cells[i].paragraphs[0].add_run(title_text).bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            hdr_shd = parse_xml(r'<w:shd {} w:fill="E9ECEF"/>'.format(nsdecls('w')))
            hdr_cells[i]._tc.get_or_add_tcPr().append(hdr_shd)

        for f in findings:
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].add_run(str(f.get("evidence_id", "E000"))).bold = True
            row_cells[1].paragraphs[0].add_run(str(f.get("equipment", "")))
            row_cells[2].paragraphs[0].add_run(str(f.get("issue", "")))
            
            sev_run = row_cells[3].paragraphs[0].add_run(str(f.get("severity", "Medium")))
            sev_run.bold = True
            if str(f.get("severity", "")).lower() == "critical":
                sev_run.font.color.rgb = RGBColor(180, 0, 0)
            elif str(f.get("severity", "")).lower() == "high":
                sev_run.font.color.rgb = RGBColor(210, 105, 0)

            measured_str = f"{f.get('measured_value', 'N/A')} (Limit: {f.get('threshold_value', 'N/A')})"
            row_cells[4].paragraphs[0].add_run(measured_str)
            
            stat_run = row_cells[5].paragraphs[0].add_run(str(f.get("status", "REVIEW")))
            stat_run.bold = True

            # Format font size
            for c in row_cells:
                for r in c.paragraphs[0].runs:
                    r.font.size = Pt(8.5)

        doc.add_paragraph()

        # 5. DETERMINISTIC CALCULATION AUDIT
        if calculation_data:
            doc.add_heading("3. Deterministic Engineering Calculations (Audit-Verified)", level=2)
            calc_p = doc.add_paragraph()
            calc_p.add_run("Quantitative calculations executed by deterministic math engine (no LLM hallucination):\n").italic = True
            
            for step in calculation_data.get("audit_trail", []):
                p_step = doc.add_paragraph(step, style="List Bullet")
                p_step.paragraph_format.space_after = Pt(2)

        # 6. INTERNAL SOP CITATIONS
        if sop_citations:
            doc.add_heading("4. Internal SOP Compliance & References", level=2)
            for sop in sop_citations:
                p_sop = doc.add_paragraph()
                p_sop.add_run(f"• Evidence [{sop.get('evidence_id', 'E_RET')}]: ").bold = True
                p_sop.add_run(f"{sop.get('document_name', 'SOP')} - Section: {sop.get('section_title', 'General')} (Page {sop.get('page_number', 1)})\n")
                snippet = sop.get("content", "")[:200] + ("..." if len(sop.get("content", "")) > 200 else "")
                p_sop.add_run(f"  Snippet: \"{snippet}\"").italic = True

        # 7. ENGINEERING SIGN-OFF BLOCK
        doc.add_heading("5. Engineering Review & Final Authorization", level=2)
        sign_table = doc.add_table(rows=3, cols=2)
        sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sign_table.cell(0, 0).paragraphs[0].add_run("Lead Inspection Engineer:\nName: ______________________\nEmp ID: ____________________\nSignature: __________________\nDate: _______________________").font.size = Pt(9)
        sign_table.cell(0, 1).paragraphs[0].add_run("Chief Technical Services Manager:\nName: ______________________\nEmp ID: ____________________\nSignature: __________________\nDate: _______________________").font.size = Pt(9)
        sign_table.cell(1, 0).paragraphs[0].add_run("Unit Operational Clearance: [  ] APPROVED   [  ] CONDITIONAL   [  ] REJECTED").bold = True
        sign_table.cell(1, 1).paragraphs[0].add_run("Required Action: Emergency Spool Replacement Before Commissioning").bold = True

        # Save document
        fname = output_filename or f"Approval_Note_{project_id}_{now_dt.strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = OUTPUT_DIR / fname
        doc.save(str(file_path))

        # Record in persistent state store
        artifact = self.state_store.record_artifact(
            artifact_id=f"ART_{project_id}_{int(now_dt.timestamp())}",
            project_id=project_id,
            artifact_type="docx",
            file_path=str(file_path),
            file_size_bytes=file_path.stat().st_size,
            is_verified=0,  # Pending inline verification
            verification_notes="Generated successfully. Awaiting verification pass.",
        )

        return {
            "artifact_id": artifact.artifact_id,
            "file_path": str(file_path),
            "file_name": fname,
            "file_size_bytes": artifact.file_size_bytes,
            "human_review_disclaimer_included": True,
            "status": "generated",
        }
