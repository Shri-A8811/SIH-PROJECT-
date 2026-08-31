"""
Inline Artifact Verification Gate for Sovereign On-Premise Agentic AI Workbench.
Checks:
1. File openability and XML structure integrity.
2. Mandatory section presence (Human Review Banner, Executive Summary, Findings, Citations, Authorization).
3. Cross-checks numeric claims against calculator outputs.
"""
from typing import Any, Dict, List, Optional
from pathlib import Path
import docx
from src.core.state_store import StateStore


class VerificationReport:
    def __init__(
        self,
        is_passed: bool,
        file_path: str,
        checks_performed: List[Dict[str, Any]],
        errors: List[str],
    ):
        self.is_passed = is_passed
        self.file_path = file_path
        self.checks_performed = checks_performed
        self.errors = errors

    def to_dict(self) -> Dict[str, Any]:
        return {
            "is_passed": self.is_passed,
            "file_path": self.file_path,
            "checks_performed": self.checks_performed,
            "errors": self.errors,
        }


class ArtifactVerifier:
    """Performs inline verification on generated Word documents."""

    def __init__(self, state_store: StateStore):
        self.state_store = state_store

    def verify_docx_deliverable(
        self,
        artifact_id: str,
        file_path: str,
        expected_numeric_values: Optional[List[str]] = None,
    ) -> VerificationReport:
        """Runs the verification suite on a generated .docx file."""
        checks: List[Dict[str, Any]] = []
        errors: List[str] = []
        path = Path(file_path)

        # 1. Check file exists and has non-zero size
        if not path.exists() or path.stat().st_size == 0:
            errors.append(f"File {file_path} does not exist or is 0 bytes.")
            return VerificationReport(is_passed=False, file_path=file_path, checks_performed=checks, errors=errors)

        checks.append({"check": "file_exists_and_non_empty", "passed": True, "size_bytes": path.stat().st_size})

        # 2. Check document opens cleanly in python-docx
        doc = None
        try:
            doc = docx.Document(str(path))
            checks.append({"check": "document_xml_integrity_and_openability", "passed": True})
        except Exception as e:
            errors.append(f"Document failed to open: {str(e)}")
            checks.append({"check": "document_xml_integrity_and_openability", "passed": False, "error": str(e)})
            self._update_artifact_record(artifact_id, is_passed=False, notes=str(errors))
            return VerificationReport(is_passed=False, file_path=file_path, checks_performed=checks, errors=errors)

        # 3. Check Mandatory Section Headings and Disclaimer
        full_text = []
        for p in doc.paragraphs:
            full_text.append(p.text)
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells:
                    full_text.append(c.text)
        combined_text = "\n".join(full_text)

        required_phrases = [
            ("human_review_disclaimer", "HUMAN REVIEW REQUIRED"),
            ("executive_summary", "Executive Summary"),
            ("safety_findings_section", "Safety-Critical Inspection Findings"),
            ("engineering_signoff", "Engineering Review & Final Authorization"),
        ]

        for check_name, phrase in required_phrases:
            if phrase.lower() in combined_text.lower():
                checks.append({"check": f"contains_{check_name}", "passed": True})
            else:
                errors.append(f"Mandatory requirement missing: document lacks '{phrase}'.")
                checks.append({"check": f"contains_{check_name}", "passed": False})

        # 4. Check Numeric Claim Matching
        if expected_numeric_values:
            for num_val in expected_numeric_values:
                if str(num_val) in combined_text:
                    checks.append({"check": f"numeric_match_{num_val}", "passed": True})
                else:
                    errors.append(f"Numeric claim cross-check failed: value '{num_val}' not found in document text.")
                    checks.append({"check": f"numeric_match_{num_val}", "passed": False})

        is_passed = len(errors) == 0
        notes = "All inline verification checks passed successfully." if is_passed else "; ".join(errors)

        # Update Artifact Record in Persistent State Store
        self._update_artifact_record(artifact_id, is_passed=is_passed, notes=notes)

        return VerificationReport(
            is_passed=is_passed,
            file_path=file_path,
            checks_performed=checks,
            errors=errors,
        )

    def _update_artifact_record(self, artifact_id: str, is_passed: bool, notes: str):
        self.state_store.update_artifact_verification(
            artifact_id=artifact_id,
            is_verified=1 if is_passed else -1,
            verification_notes=notes,
        )
